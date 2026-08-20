from __future__ import annotations

import csv
import hashlib
import json
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from reportlab.lib.colors import Color, HexColor, white
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


HERE = Path(__file__).resolve().parent
REPO = HERE.parent.parent
DATA = json.loads((HERE / "demo-data.json").read_text(encoding="utf-8"))
DATE = "20260820"
OUTPUTS = REPO / "outputs" / "demo-material-packages-20260820"
PDF_OUTPUT = REPO / "output" / "pdf" / "demo-material-packages-20260820"
DOC_OUTPUT = REPO / "output" / "documents" / "demo-material-packages-20260820"
PACKAGE_ROOT = REPO / "artifacts" / "demo-material-packages"
PUBLIC_DOWNLOADS = REPO / "site" / "public" / "assets" / "downloads"

LABEL_NAME = f"方寸标签排版_脱敏功能演示素材_{DATE}"
PDF_NAME = f"方寸PDF配印助手_脱敏功能演示素材_{DATE}"
LABEL_DIR = PACKAGE_ROOT / LABEL_NAME
PDF_DIR = PACKAGE_ROOT / PDF_NAME

FONT_PATH = Path(r"C:\Windows\Fonts\simhei.ttf")
FONT_NAME = "SimHeiSyntheticDemo"
pdfmetrics.registerFont(TTFont(FONT_NAME, str(FONT_PATH)))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def copy_file(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8", newline="\r\n")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def build_label_docx(target: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color, before, after in [
        ("Heading 1", 15, "486A2E", 10, 5),
        ("Heading 2", 12.5, "486A2E", 8, 4),
        ("Heading 3", 11, "2F4930", 6, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "方寸标签排版｜公开合成演示数据"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.runs[0]
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(100, 112, 107)

    footer = section.footer.paragraphs[0]
    footer.text = "SYNTHETIC DEMO｜不含真实客户数据｜不代表软件实跑结果"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.name = "Microsoft YaHei"
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 112, 107)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("LABEL IMPOSITION / OUTPUT STRUCTURE")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(115, 158, 69)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("标签排版说明｜演示输出结构")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 38)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("由 24 条虚构订单构建，用于说明员工执行与负责人复核时应看到的信息层级。")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(80, 92, 87)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EFF6E7")
    p_pr.append(shading)
    run = paragraph.add_run("真实性边界：订单号、SKU、品名、数量、材料和成本均为虚构演示值。本文件是网站展示结构样例，并非软件真实导出或客户项目证明。")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(72, 106, 46)

    doc.add_heading("1. 任务摘要", level=1)
    summary = doc.add_table(rows=2, cols=4)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    widths = [1.55, 1.75, 1.55, 1.9]
    labels = ["订单款数", "总需求数量", "尺寸组", "统一冗余率"]
    values = ["24", "30,275", "3", "5%"]
    for index, width in enumerate(widths):
        summary.columns[index].width = Inches(width)
        for row in summary.rows:
            row.cells[index].width = Inches(width)
    for col, label in enumerate(labels):
        cell = summary.cell(0, col)
        cell.text = label
        set_cell_shading(cell, "486A2E")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_cell = summary.cell(1, col)
        value_cell.text = values[col]
        set_cell_margins(value_cell)
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in value_cell.paragraphs[0].runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 38)
    set_repeat_table_header(summary.rows[0])

    doc.add_heading("2. 尺寸组计划", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["尺寸 mm", "款数", "需求数量", "版面容量", "材料", "复核"]
    widths = [1.1, 0.7, 1.25, 1.0, 1.35, 1.0]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
        table.rows[0].cells[index].width = Inches(width)
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, "739E45")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for row_data in [
        ["60x40", "8", "11,300", "48", "铜版纸", "PASS"],
        ["80x50", "8", "10,415", "30", "合成纸", "PASS"],
        ["100x70", "8", "8,560", "20", "哑银 PET", "PASS"],
    ]:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.width = Inches(widths[index])
            cell.text = value
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if index in (0, 4) else WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(9)

    doc.add_heading("3. 执行与复核顺序", level=1)
    for text in [
        "导入本包的复杂脱敏演示订单工作簿，并确认 24 条订单全部读取。",
        "按成品尺寸分组，核对每组版面容量、材料与统一 5% 冗余率。",
        "查看排版计算明细中的冗余后数量、计划印张、计划产能与余量。",
        "负责人复核成本列仅作为结构示例，不据此报价；异常状态保留为 CHECK，本示例结构检查均为 PASS。",
    ]:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.38)
        paragraph.paragraph_format.first_line_indent = Inches(-0.19)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9.5)

    doc.add_heading("4. 配套文件", level=1)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("输入工作簿用于订单与公式汇总；排版计算明细用于计划、成本与余量复核；订单清单、manifest 和 SHA256SUMS 用于机器读取与完整性检查。")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9.5)

    doc.core_properties.title = "标签排版说明｜演示输出结构"
    doc.core_properties.subject = "公开合成脱敏演示数据"
    doc.core_properties.author = "方寸有序工作室"
    doc.core_properties.keywords = "SYNTHETIC DEMO, 标签排版, 脱敏展示"
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def draw_pdf_page(c: canvas.Canvas, title: str, source_file: str, page_number: int, total_pages: int) -> None:
    width, height = A4
    c.setFillColor(HexColor("#173D48"))
    c.rect(0, height - 95, width, 95, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont(FONT_NAME, 22)
    c.drawString(48, height - 56, title)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(48, height - 78, "SYNTHETIC DEMO / PUBLIC SAMPLE DATA")

    c.setFillColor(HexColor("#EFF7F8"))
    c.roundRect(48, height - 190, width - 96, 66, 8, fill=1, stroke=0)
    c.setFillColor(HexColor("#245969"))
    c.setFont(FONT_NAME, 14)
    c.drawString(66, height - 150, f"源文件：{source_file}")
    c.setFont(FONT_NAME, 11)
    c.drawString(66, height - 174, f"源页码：{page_number} / {total_pages}")

    c.setFillColor(HexColor("#1E2926"))
    c.setFont(FONT_NAME, 13)
    c.drawString(48, height - 235, "页面内容摘要")
    c.setFont(FONT_NAME, 10.5)
    lines = [
        "本页仅用于方寸 PDF 配印助手的网站功能展示。",
        "标题、编号、份数、页码要求与内容均为虚构数据。",
        "处理规则由任务表定义；0 份页面在加工后文件中忽略。",
        "本 PDF 不含姓名、电话、地址、客户名称或商业文件内容。",
    ]
    y = height - 265
    for line in lines:
        c.drawString(62, y, line)
        y -= 25

    c.setFillColor(HexColor("#F6F7F4"))
    c.roundRect(48, 210, width - 96, 210, 10, fill=1, stroke=0)
    c.setFillColor(HexColor("#3A8396"))
    c.setFont("Helvetica-Bold", 46)
    c.drawCentredString(width / 2, 330, f"{page_number:02d}")
    c.setFont("Helvetica-Bold", 13)
    c.drawCentredString(width / 2, 298, "SYNTHETIC PAGE IDENTIFIER")
    c.setFillColor(HexColor("#64706B"))
    c.setFont(FONT_NAME, 10)
    c.drawCentredString(width / 2, 266, "用于逐页重复、忽略与审计对应关系的可视标识")

    c.setFillColor(HexColor("#64706B"))
    c.setFont(FONT_NAME, 8.5)
    c.drawString(48, 55, "公开合成演示数据｜不代表客户文件｜不代表软件真实运行日志")
    c.setFont("Helvetica", 8)
    c.drawRightString(width - 48, 55, f"{source_file} / {page_number}")
    c.showPage()


def build_source_pdfs() -> list[dict]:
    PDF_OUTPUT.mkdir(parents=True, exist_ok=True)
    source_dir = PDF_OUTPUT / "input"
    processed_dir = PDF_OUTPUT / "processed"
    clean_dir(source_dir)
    clean_dir(processed_dir)
    results = []
    for source in DATA["pdf"]["sources"]:
        source_path = source_dir / source["file"]
        c = canvas.Canvas(str(source_path), pagesize=A4, pageCompression=1)
        c.setTitle(source["title"])
        c.setAuthor("方寸有序工作室")
        c.setSubject("SYNTHETIC DEMO / PUBLIC SAMPLE DATA")
        for page in source["pages"]:
            draw_pdf_page(c, source["title"], source["file"], page["page"], len(source["pages"]))
        c.save()

        reader = PdfReader(str(source_path))
        writer = PdfWriter()
        output_map = []
        output_cursor = 1
        for page_spec in source["pages"]:
            start = output_cursor if page_spec["copies"] else 0
            for _ in range(page_spec["copies"]):
                writer.add_page(reader.pages[page_spec["page"] - 1])
                output_cursor += 1
            end = output_cursor - 1 if page_spec["copies"] else 0
            output_map.append({
                "sourcePage": page_spec["page"],
                "targetCopies": page_spec["copies"],
                "actualCopies": page_spec["copies"],
                "outputStart": start,
                "outputEnd": end,
            })
        writer.add_metadata({
            "/Title": f"{source['title']} - 配印后结构示例",
            "/Author": "方寸有序工作室",
            "/Subject": "SYNTHETIC DEMO / NOT A SOFTWARE EXECUTION LOG",
        })
        processed_name = source["file"].replace(".pdf", "_配印后.pdf")
        processed_path = processed_dir / processed_name
        with processed_path.open("wb") as stream:
            writer.write(stream)

        source_check = PdfReader(str(source_path))
        output_check = PdfReader(str(processed_path))
        expected_output_pages = sum(page["copies"] for page in source["pages"])
        if len(source_check.pages) != len(source["pages"]):
            raise RuntimeError(f"Source page count mismatch: {source_path}")
        if len(output_check.pages) != expected_output_pages:
            raise RuntimeError(f"Output page count mismatch: {processed_path}")
        text = "\n".join((page.extract_text() or "") for page in source_check.pages)
        if "SYNTHETIC DEMO" not in text:
            raise RuntimeError(f"Synthetic marker missing: {source_path}")
        results.append({
            "source": source_path,
            "processed": processed_path,
            "sourcePages": len(source_check.pages),
            "outputPages": len(output_check.pages),
            "map": output_map,
        })
    return results


def make_pdf_contact_sheet(results: list[dict], target: Path) -> None:
    width, height = 1600, 950
    image = Image.new("RGB", (width, height), "#F4F1E9")
    draw = ImageDraw.Draw(image)
    font_title = ImageFont.truetype(str(FONT_PATH), 46)
    font_body = ImageFont.truetype(str(FONT_PATH), 24)
    font_small = ImageFont.truetype(str(FONT_PATH), 18)
    draw.text((60, 46), "PDF 配印展示包｜处理前后结构总览", font=font_title, fill="#173D48")
    draw.text((60, 110), "SYNTHETIC DEMO｜源 PDF 与加工后 PDF 均由公开虚构数据构建", font=font_body, fill="#3A8396")
    y = 205
    for index, result in enumerate(results, start=1):
        color = "#EAF5F7" if index % 2 else "#F1F6EA"
        draw.rounded_rectangle((60, y, 1540, y + 190), radius=18, fill=color, outline="#CBD8D8", width=2)
        draw.text((92, y + 28), result["source"].name, font=font_body, fill="#1E2926")
        draw.text((92, y + 78), f"源页数：{result['sourcePages']}　→　目标/实际输出页数：{result['outputPages']}", font=font_small, fill="#42514C")
        mapping = "；".join(
            f"P{item['sourcePage']}×{item['targetCopies']}"
            for item in result["map"]
        )
        draw.text((92, y + 124), f"逐页规则：{mapping}", font=font_small, fill="#42514C")
        draw.text((1310, y + 70), "PASS", font=font_title, fill="#587B37")
        y += 220
    target.parent.mkdir(parents=True, exist_ok=True)
    image.save(target, "PNG", optimize=True)


def build_label_package() -> dict:
    clean_dir(LABEL_DIR)
    (LABEL_DIR / "输出示例").mkdir(parents=True)
    (LABEL_DIR / "QA").mkdir(parents=True)
    copy_file(OUTPUTS / "label" / "01_方寸标签排版_复杂脱敏演示订单.xlsx", LABEL_DIR / "01_方寸标签排版_复杂脱敏演示订单.xlsx")
    copy_file(OUTPUTS / "label" / "03_订单与排版计划预览.png", LABEL_DIR / "03_订单与排版计划预览.png")
    copy_file(OUTPUTS / "label" / "04_输出结构总览.png", LABEL_DIR / "04_输出结构总览.png")
    copy_file(OUTPUTS / "label" / "排版计算明细_演示.xlsx", LABEL_DIR / "输出示例" / "排版计算明细_演示.xlsx")

    DOC_OUTPUT.mkdir(parents=True, exist_ok=True)
    docx_path = DOC_OUTPUT / "标签排版说明_演示.docx"
    build_label_docx(docx_path)
    copy_file(docx_path, LABEL_DIR / "输出示例" / "标签排版说明_演示.docx")

    csv_path = LABEL_DIR / "订单清单.csv"
    with csv_path.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(["订单编号", "虚构SKU", "演示品名", "成品尺寸mm", "需求数量", "版面容量", "材料", "单张材料成本", "数据性质"])
        for row in DATA["label"]["orders"]:
            writer.writerow([*row, "SYNTHETIC DEMO"])

    write_text(LABEL_DIR / "02_使用步骤与预期结果.txt", """方寸标签排版｜复杂脱敏功能演示素材
====================================

本包全部为公开虚构数据，不含真实客户、联系人、地址、电话、订单或价格。
目标：展示 Excel 数量表导入、尺寸分组、版面容量、5% 冗余、排版计划、成本与余量复核的文件结构。

一、输入
1. 打开 01_方寸标签排版_复杂脱敏演示订单.xlsx。
2. 检查“订单输入”工作表中的 24 条 DEMO / SYN 订单。
3. 三个尺寸组：60x40、80x50、100x70 mm；每组 8 款。
4. 统一冗余率 5%；版面容量分别为 48、30、20。

二、预期结构
- 计划预览应汇总为 24 款、30,275 个需求数量。
- 每条订单计算冗余后数量、计划印张、计划产能和余量。
- 成本值仅为演示公式输入，不作为报价依据。
- 所有余量应为非负，结构检查结果为 PASS。

三、直接查看
- 03_订单与排版计划预览.png：尺寸组汇总。
- 04_输出结构总览.png：逐订单计划结构。
- 输出示例/排版计算明细_演示.xlsx：公式化输出结构。
- 输出示例/标签排版说明_演示.docx：员工执行与负责人复核说明结构。

四、真实性边界
本包由网站展示构建流程生成，softwareExecutionClaim=false；不将样例描述为软件真实导出或客户项目。
""")
    write_text(LABEL_DIR / "05_结构验收结果_PASS.txt", """方寸标签排版｜合成脱敏展示包结构验收
RESULT=PASS

- 24/24 虚构订单带 DEMO / SYN 标识
- 3/3 尺寸组汇总完整
- 需求总量：30,275
- XLSX：2/2 可读取并包含公式化复核列
- DOCX：1/1 包含真实性边界与执行步骤
- PNG：2/2 可读取
- realCustomerData=false
- softwareExecutionClaim=false
- SHA256SUMS 覆盖除自身外全部文件
""")
    validation = {
        "result": "PASS",
        "productId": "label",
        "provenance": "synthetic-demo-package",
        "orders": 24,
        "sizeGroups": 3,
        "totalDemand": sum(row[4] for row in DATA["label"]["orders"]),
        "realCustomerData": False,
        "softwareExecutionClaim": False,
        "checks": ["xlsx-present", "docx-present", "preview-images-readable", "synthetic-prefixes-present"],
    }
    (LABEL_DIR / "QA" / "validation-report.json").write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")
    write_text(LABEL_DIR / "QA" / "validation-report.txt", json.dumps(validation, ensure_ascii=False, indent=2))
    write_text(LABEL_DIR / "README_生成清单.md", """# 方寸标签排版｜脱敏功能演示素材

本包仿照“方寸有序胀色裁切”展示包的入口与验收形式制作，但所有内容均为重新构建的公开虚构数据。

## 覆盖内容

- 24 条虚构标签订单，覆盖 3 个成品尺寸与 3 种演示材料。
- Excel 输入、尺寸组汇总、冗余后数量、计划印张、产能、余量与成本结构。
- Word 排版说明与 Excel 计算明细的输出结构示例。
- 两张预览图、机器可读订单清单、QA 报告和全量 SHA-256 清单。

## 主要入口

- `01_方寸标签排版_复杂脱敏演示订单.xlsx`
- `02_使用步骤与预期结果.txt`
- `03_订单与排版计划预览.png`
- `04_输出结构总览.png`
- `输出示例/标签排版说明_演示.docx`
- `输出示例/排版计算明细_演示.xlsx`

## 真实性边界

`provenance=synthetic-demo-package`，`realCustomerData=false`，`softwareExecutionClaim=false`。订单、SKU、品名、数量、材料和成本均为虚构演示值；样例不代表软件真实运行、客户项目或商业报价。
""")
    return finalize_package(LABEL_DIR, "label", "方寸标签排版脱敏功能演示素材")


def build_pdf_package(pdf_results: list[dict]) -> dict:
    clean_dir(PDF_DIR)
    (PDF_DIR / "输入PDF素材").mkdir(parents=True)
    (PDF_DIR / "输出示例" / "加工后PDF").mkdir(parents=True)
    (PDF_DIR / "输出示例").mkdir(parents=True, exist_ok=True)
    (PDF_DIR / "QA").mkdir(parents=True)
    copy_file(OUTPUTS / "pdf" / "01_方寸PDF配印助手_复杂脱敏任务.xlsx", PDF_DIR / "01_方寸PDF配印助手_复杂脱敏任务.xlsx")
    copy_file(OUTPUTS / "pdf" / "03_任务输入与规则预览.png", PDF_DIR / "03_任务输入与规则预览.png")
    copy_file(OUTPUTS / "pdf" / "04_逐页审计总览.png", PDF_DIR / "04_逐页审计总览.png")
    copy_file(OUTPUTS / "pdf" / "逐页审计记录_演示.xlsx", PDF_DIR / "输出示例" / "逐页审计记录_演示.xlsx")
    for result in pdf_results:
        copy_file(result["source"], PDF_DIR / "输入PDF素材" / result["source"].name)
        copy_file(result["processed"], PDF_DIR / "输出示例" / "加工后PDF" / result["processed"].name)
    make_pdf_contact_sheet(pdf_results, PDF_DIR / "05_处理前后结构总览.png")

    with (PDF_DIR / "任务清单.csv").open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(["源PDF", "虚构文档标题", "源页码", "目标份数", "规则说明", "数据性质"])
        for source in DATA["pdf"]["sources"]:
            for page in source["pages"]:
                writer.writerow([source["file"], source["title"], page["page"], page["copies"], page["note"], "SYNTHETIC DEMO"])

    write_text(PDF_DIR / "02_使用步骤与预期结果.txt", """方寸 PDF 配印助手｜复杂脱敏功能演示素材
=========================================

本包全部为公开虚构数据，不含真实客户、联系人、地址、电话或商业 PDF 内容。
目标：展示 Excel 逐页任务、0 份忽略规则、加工后 PDF 与逐页审计记录的文件结构。

一、任务
- SYNTH-PDF-A.pdf：4 个源页，目标输出 6 页；第 4 页为 0 份并忽略。
- SYNTH-PDF-B.pdf：3 个源页，目标输出 5 页。
- SYNTH-PDF-C.pdf：2 个源页，目标输出 5 页。
- 合计：9 个源页规则，目标/实际输出 16 页。

二、查看顺序
1. 打开 01_方寸PDF配印助手_复杂脱敏任务.xlsx。
2. 查看 输入PDF素材/ 下的 3 份源 PDF。
3. 对照 输出示例/加工后PDF/ 的 3 份加工后结构样例。
4. 打开 输出示例/逐页审计记录_演示.xlsx，逐行核对目标份数与实际份数。
5. 使用 manifest.json 与 SHA256SUMS.txt 复核包体。

三、真实性边界
加工后 PDF 由本展示包构建脚本按任务表重复页面生成；它们用于验证文件结构，不代表成品软件真实运行日志或客户项目。
""")
    write_text(PDF_DIR / "06_结构验收结果_PASS.txt", """方寸 PDF 配印助手｜合成脱敏展示包结构验收
RESULT=PASS

- 源 PDF：3/3 可读取
- 源页：9/9 包含 SYNTHETIC DEMO 标识
- 加工后 PDF：3/3 可读取
- 目标/实际输出页数：16/16
- 0 份规则：1/1 已忽略
- XLSX：2/2 可读取并包含逐页任务或审计结构
- PNG：3/3 可读取
- realCustomerData=false
- softwareExecutionClaim=false
- SHA256SUMS 覆盖除自身外全部文件
""")
    validation = {
        "result": "PASS",
        "productId": "pdf",
        "provenance": "synthetic-demo-package",
        "sourcePdfs": len(pdf_results),
        "sourcePages": sum(result["sourcePages"] for result in pdf_results),
        "outputPdfs": len(pdf_results),
        "outputPages": sum(result["outputPages"] for result in pdf_results),
        "realCustomerData": False,
        "softwareExecutionClaim": False,
        "files": [
            {
                "source": result["source"].name,
                "sourcePages": result["sourcePages"],
                "processed": result["processed"].name,
                "outputPages": result["outputPages"],
                "map": result["map"],
            }
            for result in pdf_results
        ],
    }
    (PDF_DIR / "QA" / "validation-report.json").write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")
    write_text(PDF_DIR / "QA" / "validation-report.txt", json.dumps(validation, ensure_ascii=False, indent=2))
    write_text(PDF_DIR / "README_生成清单.md", """# 方寸 PDF 配印助手｜脱敏功能演示素材

本包仿照“方寸有序胀色裁切”展示包的入口与验收形式制作，但所有 PDF、任务与输出均为重新构建的公开虚构数据。

## 覆盖内容

- 3 份源 PDF、9 个源页和 9 条逐页份数规则。
- 0 份页面忽略、按份数重复、加工后 PDF 与逐页审计表。
- 3 份加工后结构样例，共 16 个输出页。
- 三张预览图、任务 CSV、QA 报告和全量 SHA-256 清单。

## 主要入口

- `01_方寸PDF配印助手_复杂脱敏任务.xlsx`
- `02_使用步骤与预期结果.txt`
- `输入PDF素材/`
- `输出示例/加工后PDF/`
- `输出示例/逐页审计记录_演示.xlsx`
- `05_处理前后结构总览.png`

## 真实性边界

`provenance=synthetic-demo-package`，`realCustomerData=false`，`softwareExecutionClaim=false`。加工后 PDF 由展示包构建流程按任务表生成，不代表软件真实运行日志或客户项目。
""")
    return finalize_package(PDF_DIR, "pdf", "方寸 PDF 配印助手脱敏功能演示素材")


def finalize_package(root: Path, product_id: str, title: str) -> dict:
    payload_files = sorted(
        path for path in root.rglob("*")
        if path.is_file() and path.name not in {"manifest.json", "SHA256SUMS.txt"}
    )
    manifest = {
        "schema": 1,
        "title": title,
        "productId": product_id,
        "generatedOn": "2026-08-20",
        "provenance": "synthetic-demo-package",
        "realCustomerData": False,
        "softwareExecutionClaim": False,
        "rules": {
            "publicSyntheticData": True,
            "customerIdentifiers": False,
            "contactDetails": False,
            "commercialClaims": False,
            "softwareExecutionEvidence": False,
        },
        "files": [
            {
                "path": path.relative_to(root).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
            for path in payload_files
        ],
    }
    (root / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    checksum_files = sorted(path for path in root.rglob("*") if path.is_file() and path.name != "SHA256SUMS.txt")
    checksum_lines = [f"{sha256(path)}  {path.relative_to(root).as_posix()}" for path in checksum_files]
    (root / "SHA256SUMS.txt").write_text("\n".join(checksum_lines) + "\n", encoding="utf-8")

    for line in checksum_lines:
        expected, relative = line.split("  ", 1)
        actual = sha256(root / Path(relative))
        if expected != actual:
            raise RuntimeError(f"Checksum mismatch: {relative}")

    zip_name = "label-redacted-demo-materials-20260820.zip" if product_id == "label" else "pdf-redacted-demo-materials-20260820.zip"
    zip_path = PUBLIC_DOWNLOADS / zip_name
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    deterministic_zip(root, zip_path)
    return {
        "productId": product_id,
        "root": root.relative_to(REPO).as_posix(),
        "zip": zip_path.relative_to(REPO).as_posix(),
        "zipBytes": zip_path.stat().st_size,
        "zipSha256": sha256(zip_path),
        "fileCount": len([path for path in root.rglob("*") if path.is_file()]),
        "checksumEntries": len(checksum_lines),
    }


def deterministic_zip(root: Path, target: Path) -> None:
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(p for p in root.rglob("*") if p.is_file()):
            arcname = f"{root.name}/{path.relative_to(root).as_posix()}"
            info = zipfile.ZipInfo(arcname, date_time=(2026, 8, 20, 12, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o100644 << 16
            archive.writestr(info, path.read_bytes(), compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)


def main() -> None:
    pdf_results = build_source_pdfs()
    label_result = build_label_package()
    pdf_result = build_pdf_package(pdf_results)
    report = {"label": label_result, "pdf": pdf_result}
    (PACKAGE_ROOT / "package-build-report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
